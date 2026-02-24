#!/usr/bin/env python3
"""
Sync content from LaTeX to an existing DOCX template without full conversion.

This script patches existing DOCX blocks and can insert missing blocks:
- list items under headings: Scope, Prerequisites, Maintenance Notes
- table rows under subsections:
  - Top Navigation
  - Left Navigation (Common Signed-out Items)
  - Home Feed Video Card
- numbered lists under flow subsections:
  - Flow A: Search for a Video
  - Flow B: Open a Video Watch Page
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except Exception:  # pragma: no cover - optional dependency for dynamic mode
    Document = None
    OxmlElement = None
    qn = None
    Paragraph = Any

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_XML = "http://www.w3.org/XML/1998/namespace"
NS = {"w": NS_W}


def w(tag: str) -> str:
    return f"{{{NS_W}}}{tag}"


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Sync LaTeX content into DOCX template")
    p.add_argument("--tex", type=Path, default=Path("main.tex"))
    p.add_argument("--docx", type=Path, default=Path("manual_styled_v3.docx"))
    p.add_argument("--spec", type=Path, default=None, help="Dynamic manual spec (enables token-based block sync)")
    p.add_argument("--out", type=Path, default=None, help="Output docx (default: overwrite --docx)")
    p.add_argument("--dry-run", action="store_true")
    return p.parse_args()


def canonical(text: str) -> str:
    t = re.sub(r"\\text[A-Za-z]+\{\}", "", text)
    t = re.sub(r"\\[A-Za-z]+\{([^}]*)\}", r"\1", t)
    t = re.sub(r"[^a-z0-9]+", "", t.lower())
    return t


def clean_latex_text(s: str) -> str:
    s = re.sub(r"\\texttt\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\url\{([^}]*)\}", r"\1", s)
    s = s.replace(r"\&", "&").replace(r"\%", "%").replace(r"\_", "_")
    s = s.replace(r"\$", "$").replace(r"\#", "#").replace(r"\{", "{").replace(r"\}", "}")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def parse_tex(tex_path: Path):
    tex = tex_path.read_text(encoding="utf-8")

    def get_itemize(section_name: str) -> list[str]:
        pattern = rf"\\section\{{{re.escape(section_name)}\}}.*?\\begin\{{itemize\}}(.*?)\\end\{{itemize\}}"
        m = re.search(pattern, tex, flags=re.S)
        if not m:
            return []
        body = m.group(1)
        return [clean_latex_text(x.strip()) for x in re.findall(r"\\item\s+(.*)", body)]

    def get_table(subsection_name: str) -> list[list[str]]:
        pattern = (
            rf"\\subsection\{{{re.escape(subsection_name)}\}}.*?\\begin\{{longtable\}}.*?"
            rf"\\endhead(.*?)\\bottomrule"
        )
        m = re.search(pattern, tex, flags=re.S)
        if not m:
            return []
        rows: list[list[str]] = []
        for line in m.group(1).splitlines():
            line = line.strip()
            if not line or not line.endswith(r"\\"):
                continue
            line = line[:-2].strip()
            cells = [clean_latex_text(c.strip()) for c in line.split("&")]
            if len(cells) >= 3:
                rows.append(cells[:3])
        return rows

    def get_enumerate(subsection_name: str) -> list[str]:
        pattern = rf"\\subsection\{{{re.escape(subsection_name)}\}}.*?\\begin\{{enumerate\}}(.*?)\\end\{{enumerate\}}"
        m = re.search(pattern, tex, flags=re.S)
        if not m:
            return []
        body = m.group(1)
        return [clean_latex_text(x.strip()) for x in re.findall(r"\\item\s+(.*)", body)]

    lists = {
        "Scope": get_itemize("Scope"),
        "Prerequisites": get_itemize("Prerequisites"),
        "Maintenance Notes": get_itemize("Maintenance Notes"),
    }
    tables = {
        "Top Navigation": get_table("Top Navigation"),
        "Left Navigation (Common Signed-out Items)": get_table("Left Navigation (Common Signed-out Items)"),
        "Home Feed Video Card": get_table("Home Feed Video Card"),
    }
    enums = {
        "Flow A: Search for a Video": get_enumerate("Flow A: Search for a Video"),
        "Flow B: Open a Video Watch Page": get_enumerate("Flow B: Open a Video Watch Page"),
    }
    build_lines = [
        "Run this command in the same directory:",
        "latexmk -pdf main.tex",
        "Output: main.pdf",
    ]
    return lists, tables, enums, build_lines


def paragraph_text(p: ET.Element) -> str:
    return "".join((t.text or "") for t in p.findall(".//w:t", NS)).strip()


def is_heading_paragraph(p: ET.Element) -> bool:
    style = p.find("./w:pPr/w:pStyle", NS)
    if style is None:
        return False
    val = style.attrib.get(w("val"), "")
    return val.startswith("Heading")


def is_list_paragraph(p: ET.Element) -> bool:
    return p.find("./w:pPr/w:numPr", NS) is not None


def set_paragraph_text(p: ET.Element, text: str) -> None:
    ppr = p.find("./w:pPr", NS)
    for child in list(p):
        if child.tag != w("pPr"):
            p.remove(child)
    r = ET.Element(w("r"))
    t = ET.SubElement(r, w("t"))
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        t.set(f"{{{NS_XML}}}space", "preserve")
    t.text = text
    if ppr is None:
        p.insert(0, r)
    else:
        p.append(r)


def top_level_heading_map(body: ET.Element) -> dict[str, int]:
    out: dict[str, int] = {}
    children = list(body)
    for i, el in enumerate(children):
        if el.tag != w("p"):
            continue
        if not is_heading_paragraph(el):
            continue
        txt = paragraph_text(el)
        if txt:
            out[canonical(txt)] = i
    return out


def find_heading_index(body: ET.Element, heading: str) -> int | None:
    cmap = top_level_heading_map(body)
    target = canonical(heading)
    if target in cmap:
        return cmap[target]
    for k, v in cmap.items():
        if target in k or k in target:
            return v
    return None


def find_first_list_template(body: ET.Element) -> ET.Element | None:
    for el in list(body):
        if el.tag == w("p") and is_list_paragraph(el):
            return el
    return None


def find_list_template_under_heading(body: ET.Element, heading: str) -> ET.Element | None:
    idx = find_heading_index(body, heading)
    if idx is None:
        return None
    children = list(body)
    for j in range(idx + 1, len(children)):
        el = children[j]
        if el.tag == w("p") and is_heading_paragraph(el):
            break
        if el.tag == w("p") and is_list_paragraph(el):
            return el
    return None


def make_heading(text: str, level: int) -> ET.Element:
    p = ET.Element(w("p"))
    ppr = ET.SubElement(p, w("pPr"))
    ps = ET.SubElement(ppr, w("pStyle"))
    ps.set(w("val"), f"Heading{level}")
    set_paragraph_text(p, text)
    return p


def heading_template(body: ET.Element, level: int) -> ET.Element | None:
    target_style = f"Heading{level}"
    for p in list(body):
        if p.tag != w("p"):
            continue
        st = p.find("./w:pPr/w:pStyle", NS)
        if st is None or st.attrib.get(w("val"), "") != target_style:
            continue
        # Prefer templates that already contain SectionNumber run.
        for r in p.findall("./w:r", NS):
            rs = r.find("./w:rPr/w:rStyle", NS)
            if rs is not None and rs.attrib.get(w("val"), "") == "SectionNumber":
                return p
    # fallback: any heading style paragraph
    for p in list(body):
        if p.tag != w("p"):
            continue
        st = p.find("./w:pPr/w:pStyle", NS)
        if st is not None and st.attrib.get(w("val"), "") == target_style:
            return p
    return None


def make_heading_like_template(body: ET.Element, text: str, level: int) -> ET.Element:
    tpl = heading_template(body, level)
    if tpl is None:
        return make_heading(text, level)

    p = copy.deepcopy(tpl)
    # Keep SectionNumber run and tab run; replace remaining text runs with our heading text.
    for r in list(p.findall("./w:r", NS)):
        keep = False
        rs = r.find("./w:rPr/w:rStyle", NS)
        if rs is not None and rs.attrib.get(w("val"), "") == "SectionNumber":
            keep = True
        if r.find("./w:tab", NS) is not None:
            keep = True
        if not keep:
            p.remove(r)

    new_r = ET.Element(w("r"))
    t = ET.SubElement(new_r, w("t"))
    t.text = f" {text}"
    p.append(new_r)
    return p


def clone_list_paragraph_with_text(template: ET.Element, text: str) -> ET.Element:
    p = copy.deepcopy(template)
    set_paragraph_text(p, text)
    return p


def make_list_number_paragraph(text: str, num_id: str) -> ET.Element:
    p = ET.Element(w("p"))
    ppr = ET.SubElement(p, w("pPr"))
    numpr = ET.SubElement(ppr, w("numPr"))
    ilvl = ET.SubElement(numpr, w("ilvl"))
    ilvl.set(w("val"), "0")
    numid = ET.SubElement(numpr, w("numId"))
    numid.set(w("val"), str(num_id))
    set_paragraph_text(p, text)
    return p


def prepare_decimal_num_ids(docx_path: Path) -> tuple[str, str, bytes | None]:
    default = "1003"
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            if "word/numbering.xml" not in zf.namelist():
                return default, default, None
            num_xml = zf.read("word/numbering.xml")
            num_root = ET.fromstring(num_xml)
    except Exception:
        return default, default, None

    abstract_fmt: dict[str, str] = {}
    for absn in num_root.findall(".//w:abstractNum", NS):
        aid = absn.attrib.get(w("abstractNumId"), "")
        fmt = absn.find(".//w:numFmt", NS)
        if aid and fmt is not None:
            abstract_fmt[aid] = fmt.attrib.get(w("val"), "")

    decimal_nums: list[tuple[str, str]] = []
    for num in num_root.findall(".//w:num", NS):
        nid = num.attrib.get(w("numId"), "")
        absid = num.find("./w:abstractNumId", NS)
        aid = absid.attrib.get(w("val"), "") if absid is not None else ""
        if nid and abstract_fmt.get(aid) == "decimal":
            decimal_nums.append((nid, aid))

    if not decimal_nums:
        return default, default, None

    flow_a_num, flow_a_abs = decimal_nums[0]
    if len(decimal_nums) >= 2:
        flow_b_num = decimal_nums[1][0]
        # Force restart at 1 for flow B numbering sequence.
        updated = ensure_num_start_override(num_root, flow_b_num)
        new_xml = ET.tostring(num_root, encoding="utf-8", xml_declaration=True) if updated else None
        return flow_a_num, flow_b_num, new_xml

    # Create a second decimal num id that shares same abstract decimal definition.
    existing_ids: list[int] = []
    for num in num_root.findall(".//w:num", NS):
        nid = num.attrib.get(w("numId"), "")
        if nid.isdigit():
            existing_ids.append(int(nid))
    new_id = str((max(existing_ids) + 1) if existing_ids else 2000)

    new_num = ET.Element(w("num"))
    new_num.set(w("numId"), new_id)
    abs_elem = ET.SubElement(new_num, w("abstractNumId"))
    abs_elem.set(w("val"), flow_a_abs)
    lvl_ovr = ET.SubElement(new_num, w("lvlOverride"))
    lvl_ovr.set(w("ilvl"), "0")
    start_ovr = ET.SubElement(lvl_ovr, w("startOverride"))
    start_ovr.set(w("val"), "1")
    num_root.append(new_num)

    new_xml = ET.tostring(num_root, encoding="utf-8", xml_declaration=True)
    return flow_a_num, new_id, new_xml


def ensure_num_start_override(num_root: ET.Element, num_id: str) -> bool:
    changed = False
    target = None
    for num in num_root.findall(".//w:num", NS):
        nid = num.attrib.get(w("numId"), "")
        if nid == str(num_id):
            target = num
            break
    if target is None:
        return False

    lvl = None
    for lo in target.findall("./w:lvlOverride", NS):
        if lo.attrib.get(w("ilvl"), "") == "0":
            lvl = lo
            break
    if lvl is None:
        lvl = ET.SubElement(target, w("lvlOverride"))
        lvl.set(w("ilvl"), "0")
        changed = True

    start = lvl.find("./w:startOverride", NS)
    if start is None:
        start = ET.SubElement(lvl, w("startOverride"))
        changed = True
    if start.attrib.get(w("val")) != "1":
        start.set(w("val"), "1")
        changed = True
    return changed


def table_rows(tbl: ET.Element) -> list[ET.Element]:
    return [tr for tr in tbl.findall("./w:tr", NS)]


def set_cell_text(cell: ET.Element, text: str) -> None:
    p = cell.find("./w:p", NS)
    if p is None:
        p = ET.SubElement(cell, w("p"))
    set_paragraph_text(p, text)


def build_table_from_template(template_tbl: ET.Element, rows: list[list[str]]) -> ET.Element:
    tbl = copy.deepcopy(template_tbl)
    trs = table_rows(tbl)
    if len(trs) < 2:
        return tbl
    header = trs[0]
    data_template = trs[1]
    for tr in trs[1:]:
        tbl.remove(tr)
    for row in rows:
        tr = copy.deepcopy(data_template)
        tcs = tr.findall("./w:tc", NS)
        for i in range(min(3, len(tcs))):
            set_cell_text(tcs[i], row[i] if i < len(row) else "")
        tbl.append(tr)
    _ = header
    return tbl


def build_table_with_header(template_tbl: ET.Element, header_cells: list[str], rows: list[list[str]]) -> ET.Element:
    tbl = build_table_from_template(template_tbl, rows if rows else [["", "", ""]])
    trs = table_rows(tbl)
    if not trs:
        return tbl
    hcells = trs[0].findall("./w:tc", NS)
    for i, txt in enumerate((header_cells + ["", "", ""])[:3]):
        if i < len(hcells):
            set_cell_text(hcells[i], txt)
    return tbl


def insert_after_index(body: ET.Element, idx: int, elems: list[ET.Element]) -> None:
    pos = idx + 1
    for el in elems:
        body.insert(pos, el)
        pos += 1


def section_end_index(body: ET.Element, start_idx: int) -> int:
    children = list(body)
    i = start_idx + 1
    while i < len(children):
        el = children[i]
        if el.tag == w("p") and is_heading_paragraph(el):
            return i
        i += 1
    return len(children)


def normalize_heading_manual_numbers(body: ET.Element) -> int:
    """
    Remove manually typed number prefixes in heading text runs and keep only
    Word numbering (if present).
    """
    changed = 0
    for p in list(body):
        if p.tag != w("p") or not is_heading_paragraph(p):
            continue
        saw_section_number = False
        for r in p.findall("./w:r", NS):
            rstyle = r.find("./w:rPr/w:rStyle", NS)
            is_section_number = (
                rstyle is not None and rstyle.attrib.get(w("val"), "") == "SectionNumber"
            )
            if is_section_number:
                saw_section_number = True
                continue
            t = r.find("./w:t", NS)
            if t is None or not (t.text or ""):
                continue
            raw = t.text or ""
            if not saw_section_number:
                # If heading has no auto-number run, do not strip manual numbers.
                break
            new = re.sub(r"^\s*\d+(\.\d+)*\.?\s*", "", raw)
            if new and not new.startswith((" ", "\t")):
                new = " " + new
            if new != raw:
                t.text = new
                changed += 1
            break
    return changed


def set_heading_number_and_title(p: ET.Element, level: int, number_text: str, title_text: str) -> None:
    # Ensure heading style
    ppr = p.find("./w:pPr", NS)
    if ppr is None:
        ppr = ET.SubElement(p, w("pPr"))
    pstyle = ppr.find("./w:pStyle", NS)
    if pstyle is None:
        pstyle = ET.SubElement(ppr, w("pStyle"))
    pstyle.set(w("val"), f"Heading{level}")

    # Replace runs with: [SectionNumber][tab][title]
    for r in list(p.findall("./w:r", NS)):
        p.remove(r)

    r_num = ET.SubElement(p, w("r"))
    rpr = ET.SubElement(r_num, w("rPr"))
    rs = ET.SubElement(rpr, w("rStyle"))
    rs.set(w("val"), "SectionNumber")
    t_num = ET.SubElement(r_num, w("t"))
    t_num.text = number_text

    r_tab = ET.SubElement(p, w("r"))
    ET.SubElement(r_tab, w("tab"))

    r_title = ET.SubElement(p, w("r"))
    t_title = ET.SubElement(r_title, w("t"))
    t_title.text = title_text


def enforce_heading_numbers(body: ET.Element) -> int:
    changed = 0
    targets = [
        ("Scope", 1, "1"),
        ("Prerequisites", 1, "2"),
        ("Home Page Overview", 1, "3"),
        ("Links and Buttons Mapping", 1, "4"),
        ("Top Navigation", 2, "4.1"),
        ("Left Navigation (Common Signed-out Items)", 2, "4.2"),
        ("Home Feed Video Card", 2, "4.3"),
        ("Example Task Flows", 1, "5"),
        ("Flow A: Search for a Video", 2, "5.1"),
        ("Flow B: Open a Video Watch Page", 2, "5.2"),
        ("Maintenance Notes", 1, "6"),
        ("Build", 1, "7"),
    ]

    for title, level, num in targets:
        idx = find_heading_index(body, title)
        if idx is None:
            continue
        p = list(body)[idx]
        before = paragraph_text(p)
        set_heading_number_and_title(p, level, num, title)
        after = paragraph_text(p)
        if before != after:
            changed += 1
    return changed


def enforce_section_lists_and_build(
    body: ET.Element,
    lists: dict[str, list[str]],
    enums: dict[str, list[str]],
    build_lines: list[str],
    flow_a_num_id: str,
    flow_b_num_id: str,
) -> int:
    changed = 0
    children = list(body)
    bullet_tpl = find_first_list_template(body)
    enum_tpl = find_list_template_under_heading(body, "Flow A: Search for a Video")
    if enum_tpl is None:
        enum_tpl = bullet_tpl

    def rewrite_list_section(
        heading: str,
        items: list[str],
        template: ET.Element | None,
        numbered: bool = False,
        number_num_id: str = "",
    ):
        nonlocal changed
        if template is None:
            return
        hi = find_heading_index(body, heading)
        if hi is None:
            return
        end = section_end_index(body, hi)
        cur = list(body)
        # remove existing list paragraphs in this section
        for el in cur[hi + 1 : end]:
            if el.tag == w("p") and is_list_paragraph(el):
                body.remove(el)
                changed += 1
            elif el.tag == w("p"):
                t = paragraph_text(el).strip()
                # Clean legacy plain-text numbered bullets from previous sync mode.
                if t:
                    legacy_match = False
                    for i, it in enumerate(items, start=1):
                        if t == it or t == f"{i}. {it}":
                            legacy_match = True
                            break
                    if legacy_match:
                        body.remove(el)
                        changed += 1
        # insert desired items right after heading
        insert_idx = find_heading_index(body, heading)
        if insert_idx is None:
            return
        if numbered:
            elems = []
            for t in items:
                elems.append(make_list_number_paragraph(t, number_num_id))
        else:
            elems = [clone_list_paragraph_with_text(template, t) for t in items]
        if elems:
            insert_after_index(body, insert_idx, elems)
            changed += len(elems)

    rewrite_list_section(
        "Flow A: Search for a Video",
        enums.get("Flow A: Search for a Video", []),
        enum_tpl,
        numbered=True,
        number_num_id=flow_a_num_id,
    )
    rewrite_list_section(
        "Flow B: Open a Video Watch Page",
        enums.get("Flow B: Open a Video Watch Page", []),
        enum_tpl,
        numbered=True,
        number_num_id=flow_b_num_id,
    )
    rewrite_list_section("Maintenance Notes", lists.get("Maintenance Notes", []), bullet_tpl)

    # Force Build section text to match LaTeX intent.
    build_idx = find_heading_index(body, "Build")
    if build_idx is not None:
        end = section_end_index(body, build_idx)
        cur = list(body)
        for el in cur[build_idx + 1 : end]:
            body.remove(el)
            changed += 1
        insert_idx = find_heading_index(body, "Build")
        if insert_idx is not None:
            elems: list[ET.Element] = []
            for line in build_lines:
                p = ET.Element(w("p"))
                set_paragraph_text(p, line)
                elems.append(p)
            insert_after_index(body, insert_idx, elems)
            changed += len(elems)

    return changed


def sync_list_under_heading(body: ET.Element, heading: str, items: list[str]) -> int:
    if not items:
        return 0
    idx = find_heading_index(body, heading)
    if idx is None:
        return 0

    children = list(body)
    j = idx + 1
    list_idxs: list[int] = []
    while j < len(children):
        el = children[j]
        if el.tag == w("p") and is_heading_paragraph(el):
            break
        if el.tag == w("p") and is_list_paragraph(el):
            list_idxs.append(j)
        elif list_idxs:
            break
        j += 1

    if not list_idxs:
        return 0

    changed = 0
    template = children[list_idxs[0]]

    while len(list_idxs) < len(items):
        insert_at = list_idxs[-1] + 1
        body.insert(insert_at, copy.deepcopy(template))
        children = list(body)
        list_idxs = list(range(list_idxs[0], list_idxs[0] + len(items)))

    while len(list_idxs) > len(items):
        rm_idx = list_idxs.pop()
        body.remove(children[rm_idx])
        children = list(body)

    for i, txt in enumerate(items):
        p = children[list_idxs[i]]
        if paragraph_text(p) != txt:
            set_paragraph_text(p, txt)
            changed += 1
    return changed


def sync_table_under_heading(body: ET.Element, heading: str, rows: list[list[str]]) -> int:
    if not rows:
        return 0
    idx = find_heading_index(body, heading)
    if idx is None:
        return 0

    children = list(body)
    t_idx = None
    for j in range(idx + 1, len(children)):
        if children[j].tag == w("tbl"):
            t_idx = j
            break
        if children[j].tag == w("p") and is_heading_paragraph(children[j]):
            break
    if t_idx is None:
        return 0

    table = children[t_idx]
    trs = table_rows(table)
    if len(trs) < 2:
        return 0

    data_rows = trs[1:]
    template = data_rows[0]
    changed = 0

    while len(data_rows) < len(rows):
        table.append(copy.deepcopy(template))
        data_rows = table_rows(table)[1:]

    while len(data_rows) > len(rows):
        table.remove(data_rows[-1])
        data_rows = table_rows(table)[1:]

    for r_idx, row in enumerate(rows):
        tr = data_rows[r_idx]
        tcs = tr.findall("./w:tc", NS)
        for c in range(min(3, len(tcs))):
            if paragraph_text(tcs[c]) != row[c]:
                set_cell_text(tcs[c], row[c])
                changed += 1
    return changed


def ensure_missing_blocks(body: ET.Element, lists: dict[str, list[str]], tables: dict[str, list[list[str]]], enums: dict[str, list[str]]) -> int:
    changed = 0
    list_tpl = find_first_list_template(body)
    enum_tpl = find_list_template_under_heading(body, "Flow A: Search for a Video")
    if enum_tpl is None:
        enum_tpl = list_tpl

    table_tpl = None
    for el in list(body):
        if el.tag == w("tbl"):
            table_tpl = el
            break

    # Add missing link subsections
    top_nav_idx = find_heading_index(body, "Top Navigation")
    left_idx = find_heading_index(body, "Left Navigation (Common Signed-out Items)")
    video_idx = find_heading_index(body, "Home Feed Video Card")
    if top_nav_idx is not None and table_tpl is not None:
        if left_idx is None:
            insert_after_index(
                body,
                top_nav_idx,
                [
                    make_heading_like_template(body, "Left Navigation (Common Signed-out Items)", 2),
                    build_table_from_template(table_tpl, tables.get("Left Navigation (Common Signed-out Items)", [])),
                ],
            )
            changed += 1
            left_idx = find_heading_index(body, "Left Navigation (Common Signed-out Items)")
        if video_idx is None and left_idx is not None:
            insert_after_index(
                body,
                left_idx,
                [
                    make_heading_like_template(body, "Home Feed Video Card", 2),
                    build_table_from_template(table_tpl, tables.get("Home Feed Video Card", [])),
                ],
            )
            changed += 1

    # Add missing Flow B subsection
    flow_a_idx = find_heading_index(body, "Flow A: Search for a Video")
    flow_b_idx = find_heading_index(body, "Flow B: Open a Video Watch Page")
    if flow_a_idx is not None and flow_b_idx is None and enum_tpl is not None:
        elems = [make_heading_like_template(body, "Flow B: Open a Video Watch Page", 2)]
        for item in enums.get("Flow B: Open a Video Watch Page", []):
            elems.append(clone_list_paragraph_with_text(enum_tpl, item))
        insert_after_index(body, flow_a_idx, elems)
        changed += 1

    # Add missing Maintenance Notes section
    maint_idx = find_heading_index(body, "Maintenance Notes")
    ex_idx = find_heading_index(body, "Example Task Flows")
    flow_b_idx = find_heading_index(body, "Flow B: Open a Video Watch Page")
    flow_a_idx = find_heading_index(body, "Flow A: Search for a Video")
    maint_anchor = flow_b_idx if flow_b_idx is not None else (flow_a_idx if flow_a_idx is not None else ex_idx)
    if maint_idx is None and maint_anchor is not None and list_tpl is not None:
        elems = [make_heading_like_template(body, "Maintenance Notes", 1)]
        for item in lists.get("Maintenance Notes", []):
            elems.append(clone_list_paragraph_with_text(list_tpl, item))
        insert_after_index(body, maint_anchor, elems)
        changed += 1

    # Add missing Build section
    build_idx = find_heading_index(body, "Build")
    maint_idx = find_heading_index(body, "Maintenance Notes")
    if build_idx is None:
        anchor = maint_idx if maint_idx is not None else maint_anchor
        if anchor is not None:
            p = ET.Element(w("p"))
            set_paragraph_text(p, "Run this command in the same directory: latexmk -pdf main.tex")
            insert_after_index(body, anchor, [make_heading_like_template(body, "Build", 1), p])
            changed += 1

    return changed


def rebuild_links_tables(body: ET.Element, tables: dict[str, list[list[str]]]) -> int:
    changed = 0
    links_idx = find_heading_index(body, "Links and Buttons Mapping")
    example_idx = find_heading_index(body, "Example Task Flows")
    if links_idx is None or example_idx is None:
        return 0

    children = list(body)
    template_tbl = None
    # Prefer a table already in this section as template.
    for el in children[links_idx + 1 : example_idx]:
        if el.tag == w("tbl"):
            template_tbl = el
            break
    # Fallback to first table in doc.
    if template_tbl is None:
        for el in children:
            if el.tag == w("tbl"):
                template_tbl = el
                break
    if template_tbl is None:
        return 0

    # Remove all current tables in this section.
    for el in children[links_idx + 1 : example_idx]:
        if el.tag == w("tbl"):
            body.remove(el)
            changed += 1

    specs = [
        ("Top Navigation", ["Control", "Type", "Function"], tables.get("Top Navigation", [])),
        ("Left Navigation (Common Signed-out Items)", ["Item", "Type", "Function"], tables.get("Left Navigation (Common Signed-out Items)", [])),
        ("Home Feed Video Card", ["Area", "Type", "Function"], tables.get("Home Feed Video Card", [])),
    ]

    # Insert in reverse to preserve heading indexes.
    for heading, headers, rows in reversed(specs):
        hi = find_heading_index(body, heading)
        if hi is None:
            continue
        tbl = build_table_with_header(template_tbl, headers, rows)
        body.insert(hi + 1, tbl)
        changed += 1

    return changed


def write_docx_with_updated_document_xml(
    src_docx: Path,
    out_docx: Path,
    new_doc_xml: bytes,
    new_numbering_xml: bytes | None = None,
) -> None:
    with zipfile.ZipFile(src_docx, "r") as zin:
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = new_doc_xml
                elif item.filename == "word/numbering.xml" and new_numbering_xml is not None:
                    data = new_numbering_xml
                zout.writestr(item, data)
    out_docx.write_bytes(tmp_path.read_bytes())
    tmp_path.unlink(missing_ok=True)


def remove_docx_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def insert_paragraph_after_docx(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text is not None:
        p.add_run(text)
    return p


def max_attr_int(elements: list[Any], attr_name: str, default: int) -> int:
    vals: list[int] = []
    for el in elements:
        v = el.get(qn(attr_name), "")
        if v.isdigit():
            vals.append(int(v))
    return max(vals) if vals else default


def ensure_num_id_for_format(doc: Document, num_fmt: str, restart: bool = False) -> str:
    numbering = doc.part.numbering_part.element
    abstract_nums = list(numbering.findall(qn("w:abstractNum")))
    nums = list(numbering.findall(qn("w:num")))

    abstract_meta: dict[str, dict[str, str]] = {}
    for absn in abstract_nums:
        aid = absn.get(qn("w:abstractNumId"), "")
        lvl0 = None
        for lvl in absn.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl"), "") == "0":
                lvl0 = lvl
                break
        if aid and lvl0 is not None:
            fmt = lvl0.find(qn("w:numFmt"))
            lvl_text = lvl0.find(qn("w:lvlText"))
            if fmt is not None:
                abstract_meta[aid] = {
                    "fmt": fmt.get(qn("w:val"), ""),
                    "lvl_text": (lvl_text.get(qn("w:val"), "") if lvl_text is not None else ""),
                }

    target_abs_id = ""
    for aid, meta in abstract_meta.items():
        if meta.get("fmt") != num_fmt:
            continue
        if num_fmt == "bullet":
            # Some generators produce blank bullet glyph definitions (invisible markers).
            # Only reuse bullet definitions with a visible lvlText.
            if not meta.get("lvl_text", "").strip():
                continue
        target_abs_id = aid
        break

    if not target_abs_id:
        new_abs_id = str(max_attr_int(abstract_nums, "w:abstractNumId", 1999) + 1)
        absn = OxmlElement("w:abstractNum")
        absn.set(qn("w:abstractNumId"), new_abs_id)

        mlt = OxmlElement("w:multiLevelType")
        mlt.set(qn("w:val"), "singleLevel")
        absn.append(mlt)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "%1." if num_fmt == "decimal" else "•")
        lvl.append(lvl_text)
        if num_fmt == "bullet":
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), "Symbol")
            rfonts.set(qn("w:hAnsi"), "Symbol")
            rpr.append(rfonts)
            lvl.append(rpr)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        absn.append(lvl)
        numbering.append(absn)
        target_abs_id = new_abs_id
        nums = list(numbering.findall(qn("w:num")))

    if not restart:
        for num in nums:
            nid = num.get(qn("w:numId"), "")
            absid = num.find(qn("w:abstractNumId"))
            if nid and absid is not None and absid.get(qn("w:val"), "") == target_abs_id:
                return nid

    new_num_id = str(max_attr_int(nums, "w:numId", 999) + 1)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), new_num_id)
    absid = OxmlElement("w:abstractNumId")
    absid.set(qn("w:val"), target_abs_id)
    num.append(absid)

    if restart and num_fmt == "decimal":
        lvl_ovr = OxmlElement("w:lvlOverride")
        lvl_ovr.set(qn("w:ilvl"), "0")
        start_ovr = OxmlElement("w:startOverride")
        start_ovr.set(qn("w:val"), "1")
        lvl_ovr.append(start_ovr)
        num.append(lvl_ovr)

    numbering.append(num)
    return new_num_id


def set_paragraph_numpr(paragraph: Paragraph, num_id: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:numPr"):
            ppr.remove(child)

    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)


def sync_dynamic_from_spec(args: argparse.Namespace) -> int:
    if args.spec is None:
        return 1
    if Document is None or OxmlElement is None or qn is None:
        raise SystemExit("python-docx is required for --spec dynamic sync mode")

    spec = json.loads(args.spec.read_text(encoding="utf-8"))
    block_map: dict[str, dict] = {}
    for section in spec.get("sections", []):
        for block in section.get("blocks", []):
            bid = block.get("block_id")
            if bid:
                block_map[bid] = block

    block_token_re = re.compile(r"(?:\[\[)?MANUAL_BLOCK:([A-Za-z0-9_.:-]+)(?:\]\])?")
    fig_token_re = re.compile(r"(?:\[\[)?MANUAL_FIG:([A-Za-z0-9_.:-]+)(?:\]\])?")

    doc = Document(str(args.docx))
    bullet_num_id = ensure_num_id_for_format(doc, "bullet", restart=False)
    changes = 0
    skipped = 0

    for p in list(doc.paragraphs):
        text = (p.text or "").strip()
        if fig_token_re.fullmatch(text):
            continue
        m = block_token_re.fullmatch(text)
        if not m:
            continue

        block_id = m.group(1).strip()
        block = block_map.get(block_id)
        if not block:
            p.text = ""
            skipped += 1
            changes += 1
            continue

        btype = block.get("type")
        if btype == "paragraph":
            p.text = block.get("text", "")
            changes += 1
            continue

        if btype in ("bullet_list", "numbered_list"):
            items = list(block.get("items", []))
            if not items:
                p.text = ""
                changes += 1
                continue
            num_id = bullet_num_id
            if btype == "numbered_list":
                # New numId per ordered block to force restart at 1.
                num_id = ensure_num_id_for_format(doc, "decimal", restart=True)

            p.text = items[0]
            set_paragraph_numpr(p, num_id)

            tail = p
            for item in items[1:]:
                tail = insert_paragraph_after_docx(
                    tail,
                    item,
                )
                set_paragraph_numpr(tail, num_id)
                changes += 1
            changes += 1
            continue

        if btype == "table":
            cols = list(block.get("columns", []))
            rows = list(block.get("rows", []))
            if not cols:
                cols = ["Column 1", "Column 2", "Column 3"]
            table = doc.add_table(rows=1, cols=len(cols))
            for idx, c in enumerate(cols):
                table.rows[0].cells[idx].text = str(c)
            for row in rows:
                r = table.add_row().cells
                for idx, c in enumerate(cols):
                    r[idx].text = str(row[idx] if idx < len(row) else "")
            p._p.addnext(table._tbl)
            remove_docx_paragraph(p)
            changes += 1
            continue

        if btype == "figure":
            # Figure image/caption placement is handled by sync_latex_images_to_docx.py.
            p.text = ""
            changes += 1
            continue

        p.text = ""
        skipped += 1
        changes += 1

    # Cleanup unresolved block tokens to avoid leaking control markers.
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if block_token_re.search(t):
            p.text = ""
            changes += 1

    out_docx = args.out if args.out else args.docx
    print(f"docx: {args.docx}")
    print(f"spec: {args.spec}")
    print(f"out: {out_docx}")
    print(f"changes: {changes}")
    print(f"skipped_blocks: {skipped}")

    if not args.dry_run:
        doc.save(str(out_docx))
        print("updated docx")
    return 0


def legacy_main(args: argparse.Namespace) -> int:
    out_docx = args.out if args.out else args.docx

    lists, tables, enums, build_lines = parse_tex(args.tex)
    flow_a_num_id, flow_b_num_id, new_numbering_xml = prepare_decimal_num_ids(args.docx)

    with zipfile.ZipFile(args.docx, "r") as zf:
        doc_xml = zf.read("word/document.xml")
    root = ET.fromstring(doc_xml)
    body = root.find("w:body", NS)
    if body is None:
        raise SystemExit("DOCX body not found")

    changed = 0
    changed += sync_list_under_heading(body, "Scope", lists.get("Scope", []))
    changed += sync_list_under_heading(body, "Prerequisites", lists.get("Prerequisites", []))
    changed += sync_list_under_heading(body, "Maintenance Notes", lists.get("Maintenance Notes", []))

    changed += sync_table_under_heading(body, "Top Navigation", tables.get("Top Navigation", []))
    changed += sync_table_under_heading(body, "Left Navigation (Common Signed-out Items)", tables.get("Left Navigation (Common Signed-out Items)", []))
    changed += sync_table_under_heading(body, "Home Feed Video Card", tables.get("Home Feed Video Card", []))

    changed += sync_list_under_heading(body, "Flow A: Search for a Video", enums.get("Flow A: Search for a Video", []))
    changed += sync_list_under_heading(body, "Flow B: Open a Video Watch Page", enums.get("Flow B: Open a Video Watch Page", []))

    changed += ensure_missing_blocks(body, lists, tables, enums)
    changed += rebuild_links_tables(body, tables)
    changed += enforce_section_lists_and_build(
        body,
        lists,
        enums,
        build_lines,
        flow_a_num_id,
        flow_b_num_id,
    )
    changed += normalize_heading_manual_numbers(body)
    changed += enforce_heading_numbers(body)

    print(f"docx: {args.docx}")
    print(f"tex: {args.tex}")
    print(f"out: {out_docx}")
    print(f"changes: {changed}")

    if not args.dry_run and changed > 0:
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        write_docx_with_updated_document_xml(args.docx, out_docx, new_xml, new_numbering_xml)
        print("updated docx")

    return 0


def main() -> int:
    args = parse_args()
    if args.spec is not None:
        return sync_dynamic_from_spec(args)
    return legacy_main(args)


if __name__ == "__main__":
    raise SystemExit(main())
