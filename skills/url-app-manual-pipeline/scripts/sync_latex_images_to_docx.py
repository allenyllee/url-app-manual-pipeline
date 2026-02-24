#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


@dataclass
class Shot:
    img_rel: str
    caption: str
    anchor: str
    number: int


def canonical(text: str) -> str:
    text = re.sub(r"^\d+(\.\d+)*\.?\s*", "", text.strip())
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def parse_args():
    p = argparse.ArgumentParser(description="Sync screenshot placement from LaTeX to existing DOCX.")
    p.add_argument("--tex", type=Path, default=Path("main.tex"))
    p.add_argument("--docx", type=Path, default=Path("manual_styled_v3.docx"))
    p.add_argument("--spec", type=Path, default=None, help="Dynamic manual spec (enables token-based figure sync)")
    p.add_argument("--out", type=Path, default=None)
    return p.parse_args()


def parse_latex_shots(tex_path: Path) -> list[Shot]:
    tex = tex_path.read_text(encoding="utf-8")
    pattern = re.compile(r"\\screenshotbox\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}")
    rows = [(m.group(1).strip(), m.group(2).strip()) for m in pattern.finditer(tex)]

    anchors = [
        "Home Page Overview",
        "Top Navigation",
        "Left Navigation (Common Signed-out Items)",
        "Home Feed Video Card",
        "Flow A: Search for a Video",
        "Flow A: Search for a Video",
        "Flow B: Open a Video Watch Page",
        "Flow B: Open a Video Watch Page",
    ]

    shots: list[Shot] = []
    for i, (img_rel, cap) in enumerate(rows):
        anchor = anchors[i] if i < len(anchors) else "Example Task Flows"
        shots.append(Shot(img_rel=img_rel, caption=cap, anchor=anchor, number=i + 1))
    return shots


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text is not None:
        p.add_run(text)
    return p


def get_heading_level(paragraph: Paragraph) -> int | None:
    if not paragraph.style:
        return None
    name = paragraph.style.name or ""
    m = re.match(r"Heading\s+(\d+)", name)
    if not m:
        return None
    return int(m.group(1))


def find_heading_indices(doc: Document) -> dict[str, int]:
    idx: dict[str, int] = {}
    for i, p in enumerate(doc.paragraphs):
        lvl = get_heading_level(p)
        if lvl is None:
            continue
        c = canonical(p.text)
        if c:
            idx[c] = i
    return idx


def find_heading_index_fuzzy(heading_idx: dict[str, int], target: str) -> int | None:
    if target in heading_idx:
        return heading_idx[target]
    for k, v in heading_idx.items():
        if target in k or k in target:
            return v
    return None


def find_section_end_index(doc: Document, start_idx: int, start_level: int) -> int:
    end = len(doc.paragraphs) - 1
    for j in range(start_idx + 1, len(doc.paragraphs)):
        lvl = get_heading_level(doc.paragraphs[j])
        if lvl is None:
            continue
        if lvl <= start_level:
            end = j - 1
            break
    return max(end, start_idx)


def clear_existing_shot_blocks(doc: Document, captions: set[str]) -> int:
    removed = 0
    i = 0
    legacy_hints = (
        "overview",
        "navigation controls",
        "left navigation panel",
        "interactive areas",
        "flow a-",
        "flow b-",
    )
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        txt = (p.text or "").strip()
        style_name = (p.style.name or "") if p.style else ""
        is_caption_style = "caption" in style_name.lower() or style_name in ("ImageCaption", "Caption")
        is_known_caption = txt in captions or any(h in txt.lower() for h in legacy_hints)
        if is_caption_style and is_known_caption:
            remove_paragraph(p)
            removed += 1
            if i - 1 >= 0:
                prev = doc.paragraphs[i - 1]
                if paragraph_has_drawing(prev) and not (prev.text or "").strip():
                    remove_paragraph(prev)
                    removed += 1
                    i -= 1
            continue
        i += 1
    return removed


def sync_dynamic(args: argparse.Namespace) -> int:
    if args.spec is None:
        return 1

    spec = json.loads(args.spec.read_text(encoding="utf-8"))
    figure_map: dict[str, dict] = {}
    for section in spec.get("sections", []):
        for block in section.get("blocks", []):
            if block.get("type") == "figure" and block.get("figure_id"):
                figure_map[block["figure_id"]] = block

    token_re = re.compile(r"(?:\[\[)?MANUAL_FIG:([A-Za-z0-9_.:-]+)(?:\]\])?")
    doc = Document(str(args.docx))
    replaced = 0

    for p in list(doc.paragraphs):
        m = token_re.search((p.text or "").strip())
        if not m:
            continue
        fig_id = m.group(1).strip()
        fig = figure_map.get(fig_id)
        if not fig:
            p.text = ""
            continue

        img_path = (args.tex.parent / fig.get("image_rel", "")).resolve()
        if not img_path.exists():
            p.text = ""
            continue

        clear_paragraph(p)
        p.alignment = 1
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.2))

        cap_para = insert_paragraph_after(p, f"Figure {replaced + 1}. {fig.get('caption', '')}")
        cap_para.alignment = 1
        for style_name in ("ImageCaption", "Caption"):
            try:
                cap_para.style = style_name
                break
            except Exception:
                continue
        replaced += 1

    out = args.out or args.docx
    doc.save(str(out))
    print(f"docx: {args.docx}")
    print(f"spec: {args.spec}")
    print(f"out: {out}")
    print(f"inserted_blocks: {replaced}")
    return 0


def sync_legacy(args: argparse.Namespace) -> int:
    out = args.out or args.docx
    shots = parse_latex_shots(args.tex)
    doc = Document(str(args.docx))

    captions = {s.caption for s in shots} | {f"Figure {s.number}. {s.caption}" for s in shots}
    removed = clear_existing_shot_blocks(doc, captions)

    heading_idx = find_heading_indices(doc)
    inserted = 0
    tail_by_anchor: dict[str, Paragraph] = {}

    for shot in shots:
        heading_idx = find_heading_indices(doc)
        img_path = (args.tex.parent / shot.img_rel).resolve()
        if not img_path.exists():
            continue

        key = canonical(shot.anchor)
        start_idx = find_heading_index_fuzzy(heading_idx, key)
        if start_idx is None:
            continue

        start_para = doc.paragraphs[start_idx]
        level = get_heading_level(start_para) or 1

        if shot.anchor in tail_by_anchor:
            insert_after = tail_by_anchor[shot.anchor]
        else:
            end_idx = find_section_end_index(doc, start_idx, level)
            insert_after = doc.paragraphs[end_idx]

        img_para = insert_paragraph_after(insert_after)
        img_para.alignment = 1
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(6.2))

        cap_para = insert_paragraph_after(img_para, f"Figure {shot.number}. {shot.caption}")
        cap_para.alignment = 1
        for style_name in ("ImageCaption", "Caption"):
            try:
                cap_para.style = style_name
                break
            except Exception:
                continue

        tail_by_anchor[shot.anchor] = cap_para
        inserted += 1

    doc.save(str(out))
    print(f"docx: {args.docx}")
    print(f"out: {out}")
    print(f"removed_old_blocks: {removed}")
    print(f"inserted_blocks: {inserted}")
    return 0


def main() -> int:
    args = parse_args()
    if args.spec is not None:
        return sync_dynamic(args)
    return sync_legacy(args)


if __name__ == "__main__":
    raise SystemExit(main())
